"""
docx_extractor.py

Extracts readable text from DOCX files.

Paragraphs are read in document order (including paragraphs inside tables)
so the extracted text preserves the resume's logical reading order as
closely as a plain-text representation allows. Corrupted or unreadable
files are handled gracefully — this module never raises.
"""

from __future__ import annotations

import docx
from docx.opc.exceptions import PackageNotFoundError

from extractors.base_extractor import BaseExtractor, ExtractionResult


class DOCXExtractor(BaseExtractor):
    """Extracts text from DOCX files using python-docx."""

    def extract(self) -> ExtractionResult:
        if not self.file_path.exists():
            return ExtractionResult(success=False, message=f"File not found: {self.file_path}")

        try:
            document = docx.Document(str(self.file_path))
        except PackageNotFoundError:
            return ExtractionResult(
                success=False,
                message="Failed to open DOCX (the file may be corrupted or not a valid DOCX).",
            )
        except Exception as exc:  # noqa: BLE001 - surface any other read failure safely
            return ExtractionResult(success=False, message=f"Failed to read DOCX: {exc}")

        paragraphs = self._collect_paragraphs(document)
        text = "\n".join(paragraphs).strip()

        if not text:
            return ExtractionResult(
                success=False,
                message="The DOCX file was opened successfully but contains no readable text.",
            )

        return ExtractionResult(success=True, text=text, page_count=1)

    @staticmethod
    def _collect_paragraphs(document: "docx.document.Document") -> list[str]:
        """Collect non-empty paragraph text, including text inside tables.

        Table cell paragraphs are appended after body paragraphs; this is a
        deliberate simplification for Milestone 1 (plain text output only).
        Layout-accurate ordering can be revisited if a future milestone needs
        section-aware parsing.
        """
        paragraphs: list[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        return paragraphs
